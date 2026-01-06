# %%
import json
import re
import pandas as pd
import re
import string
from bertopic import BERTopic
from sentence_transformers import SentenceTransformer
import openai
from bertopic.representation import OpenAI
from docx import Document
doc = Document("H0 en.docx")
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
import nbformat
from transformers import AutoTokenizer
from umap import UMAP


tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-mpnet-base-v2")

# %%

client = openai.OpenAI(api_key="-insert-your-secrete-key")
representation_model = OpenAI(client, model = "gpt-4o-mini", chat = True)


# %%
def preprocess_text(text):
    text = re.sub(r'^[A-Za-z0-9]+[:]', '', text).strip()  
    text = re.sub(r'H0(?:-\d+)?', '', text)  
    text = re.sub(r"\bwasn't\b", "was not", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdoesn't\b", "does not", text, flags=re.IGNORECASE)
    text = re.sub(r"\bcan't\b", "cannot", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdidn't\b", "did not", text, flags=re.IGNORECASE)
    text = re.sub(r'\d+', '', text) 
    text = re.sub(r"\bwon't\b", "will not", text, flags=re.IGNORECASE)
    text = re.sub(r"\bhasn't\b", "has not", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdon't\b", "do not", text, flags=re.IGNORECASE)
    text = re.sub(r"\bcouldn't\b", "could not", text, flags=re.IGNORECASE)
    text = re.sub(r"\bthey're\b", "they are", text, flags=re.IGNORECASE)
    text = re.sub(r"\btheyre\b", "they are", text, flags=re.IGNORECASE)
    text = re.sub(r"\shes\b", "she is", text, flags=re.IGNORECASE)
    text = re.sub(r"\bim\b", "i am", text, flags=re.IGNORECASE)
    text = re.sub(r"\bi'm\b", "i am", text, flags=re.IGNORECASE)


    return text

doc = Document("H0 en.docx")
cleaned_lines = [preprocess_text(para.text) for para in doc.paragraphs if para.text.strip()]
full_text = " ".join(cleaned_lines)
texts_list = [full_text] 


# %%
def split_text_into_sentences(text):
    return re.split(r'(?<=[.!?])\s+', text)

def chunk_sentences(sentences, chunk_size):
    for i in range(0, len(sentences), chunk_size):
        yield ' '.join(sentences[i:i + chunk_size])

def split_texts_list(texts_list, chunk_size):
    result = []
    for text in texts_list:
        sentences = split_text_into_sentences(text)
        result.extend(chunk_sentences(sentences, chunk_size))
    return [chunk for chunk in result if chunk.strip()]

chunk_size = 6
split_texts = split_texts_list(texts_list, chunk_size=chunk_size)
print(f"Created {len(split_texts)} chunks.")


# %%

with open("H0_split_chunks.txt", "w", encoding="utf-8") as f:
    for i, chunk in enumerate(split_texts):
        f.write(f"[Chunk {i}]\n")
        f.write(chunk.strip() + "\n\n")

# %%
from hdbscan import HDBSCAN
custom_additions = {"yeah", "ll", "just", "yak", "um", "uh", "okay", "like", "know", "won", "wasn", "yes"
                    , "say", "did", "didn","think", "really", "hadn", "I", "like",
                    "ve", "re", "m", "t", "d", "s", "nt", "phone", "rings", "shouldn", "oh", "got", "said", "says", "shit"}
combined_stopwords = list(ENGLISH_STOP_WORDS.union(custom_additions))



vectorizer_model = TfidfVectorizer(min_df = 2, stop_words=combined_stopwords, sublinear_tf=True, ngram_range=(1, 2))

hdbscan_model = HDBSCAN(
    min_cluster_size= 4,
    min_samples = 2,
    metric = 'euclidean',
    cluster_selection_method="leaf"

)
umap_model = UMAP(n_neighbors= 8, min_dist=0.0, n_components = 10, metric='cosine', random_state = 42)

embedding_model = SentenceTransformer("medicalai/ClinicalBERT")

topic_model = BERTopic(embedding_model=embedding_model, vectorizer_model=vectorizer_model,
                         min_topic_size= 10, top_n_words= 15, hdbscan_model= hdbscan_model, umap_model=umap_model)


topics, probs = topic_model.fit_transform(split_texts)
topic_info = topic_model.get_topic_info()


# %%
print(f"Number of chunks: {len(split_texts)}")

print(topic_info)

# %%

def generate_topic_label(topic_keywords, topic_docs):

    doc_snippets = '\n'.join(f"- {doc.strip()}" for doc in topic_docs)

    prompt = (
    "You are an AI that labels discussion topics, from a cancer storytelling interview, "
    "for a software that allows doctors to browse through medical files without the need to read them from start "
    "to finish. Given the following keywords and sample documents, provide a clear and specific "
    "topic label, with enough context to be interpretable and to bring attention to the topic at hand, focusing mainly on the keyword list and using the document snippets as supporting"
    "context rather than a baseline. Max 15 words. You can hit the max. Only type the topic label and nothing else:\n\n"
    f"Keywords: {', '.join(topic_keywords)}\n\n"
    f"Sample Documents:\n{doc_snippets}\n\n"
    "Label:"
)
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        max_tokens=20,
        n=1,
        stop=None,
        temperature=0.5
    )
    label = response.choices[0].message.content.strip()
    return label

rep_docs_dict = topic_model.get_representative_docs()

topic_labels = {}
for topic_id in topic_info['Topic']:
    topic_keywords = [word[0] for word in topic_model.get_topic(topic_id)[:10]]
    topic_docs = rep_docs_dict.get(topic_id, [])
    topic_label = generate_topic_label(topic_keywords, topic_docs)
    topic_labels[topic_id] = topic_label

for topic_id, label in topic_labels.items():
    print(f"Topic {topic_id}: {label}, keywords: {topic_model.get_topic(topic_id)}")


for topic_id in topic_info['Topic']:
    keywords = [word[0] for word in topic_model.get_topic(topic_id)]
    label = topic_labels.get(topic_id, f"Topic {topic_id}")
    print(f"Topic {topic_id}: {label}")
    print(f"Top Keywords: {', '.join(keywords)}\n")

topic_model.set_topic_labels(topic_labels)


# %%

topic_id = 6
docs = topic_model.get_representative_docs(topic_id)
df = pd.DataFrame(docs, columns=["Representative Document"])
df.style.set_properties(**{'white-space': 'pre-wrap'})


# %%
topic_model.visualize_topics()
topic_model.visualize_hierarchy()

# %% all interviews

import os
from glob import glob
from hdbscan import HDBSCAN

current_dir = os.path.dirname(os.path.abspath(__file__))
interview_folder = current_dir

docx_files = glob(os.path.join(interview_folder, "*.docx"))

all_cleaned_texts = []

for file_path in docx_files:
    doc = Document(file_path)
    cleaned_lines = [preprocess_text(p.text) for p in doc.paragraphs if p.text.strip()]
    full_text = " ".join(cleaned_lines)
    all_cleaned_texts.append(full_text)

with open("merged_interviews_cleaned.txt", "w", encoding="utf-8") as f:
    for i, text in enumerate(all_cleaned_texts):
        f.write(f"[Interview {i + 1}]\n{text.strip()}\n\n")
print("Saved merged cleaned text to merged_interviews_cleaned.txt")

chunk_size = 7

def split_text_into_sentences(text):
    return re.split(r'(?<=[.!?])\s+', text)

def chunk_sentences(sentences, chunk_size):
    for i in range(0, len(sentences), chunk_size):
        yield ' '.join(sentences[i:i + chunk_size])

def split_texts_list(texts_list, chunk_size):
    result = []
    for idx, text in enumerate(texts_list):
        interview_id = f"Interview {idx + 1}"
        sentences = split_text_into_sentences(text)
        chunks = chunk_sentences(sentences, chunk_size)
        for chunk in chunks:
            if chunk.strip():
                result.append((interview_id, chunk))
    return result


split_all_chunks = split_texts_list(all_cleaned_texts, chunk_size=chunk_size)

chunk_metadata  =  [x[0] for x in split_all_chunks]
chunks_only = [x[1] for x in split_all_chunks]

print(f"Loaded {len(docx_files)} files and created {len(split_all_chunks)} chunks.")


from hdbscan import HDBSCAN
custom_additions = {"yeah", "ll", "just", "yak", "um", "uh", "okay", "like", "know", "won", "wasn", "yes"
                    , "say", "did", "didn","think", "really", "hadn", "I", "like",
                    "ve", "re", "m", "t", "d", "s", "nt", "phone", "rings", "shouldn", "oh", "got", "said", "says", "shit", "brotherinlaw",
                 "mdl", "pancreas", "pancreatic", "cancer", "hehe", "coffee"}
combined_stopwords = list(ENGLISH_STOP_WORDS.union(custom_additions))



vectorizer_model = TfidfVectorizer(min_df = 3, stop_words=combined_stopwords, sublinear_tf=True, ngram_range=(1, 2))

hdbscan_model = HDBSCAN(
    min_cluster_size= 11,
    prediction_data= True,
    cluster_selection_method= 'eom',


)
umap_model = UMAP(n_neighbors= 16, min_dist=0.2, n_components = 4, metric='cosine', random_state = 42)
embedding_model = SentenceTransformer("emilyalsentzer/Bio_ClinicalBERT")

topic_model = BERTopic(vectorizer_model=vectorizer_model,
                         min_topic_size= 10, top_n_words= 15, umap_model=umap_model, hdbscan_model= hdbscan_model, calculate_probabilities= True)


topics, probs = topic_model.fit_transform(chunks_only)
topic_info = topic_model.get_topic_info()

topics_all, probs_all = topic_model.fit_transform(chunks_only)


topic_info_all = topic_model.get_topic_info()
print(topic_info_all)


# %%

def generate_topic_label(topic_keywords, topic_docs):

    doc_snippets = '\n'.join(f"- {doc.strip()}" for doc in topic_docs)

    prompt = (
    "You are an AI that labels discussion topics, from a collection of cancer storytelling interviews, "
    "for a software that allows doctors to browse through medical files without the need to read them from start "
    "to finish. Given the following keywords, provide a clear and specific "
    "topic label, with enough context to be interpretable, focusing on the keyword list."
    "Be general. Keep it SHORT. Only type the topic label and nothing else:\n\n"
    f"Keywords: {', '.join(topic_keywords)}\n\n"
    
)
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        max_tokens=20,
        n=1,
        stop=None,
        temperature=0.5
    )
    label = response.choices[0].message.content.strip()
    return label

rep_docs_dict = topic_model.get_representative_docs()

topic_labels = {}
for topic_id in topic_info_all['Topic']:
    topic_keywords = [word[0] for word in topic_model.get_topic(topic_id)[:10]]
    topic_docs = rep_docs_dict.get(topic_id, [])
    topic_label = generate_topic_label(topic_keywords, topic_docs)
    topic_labels[topic_id] = topic_label

for topic_id, label in topic_labels.items():
    print(f"Topic {topic_id}: {label}, keywords: {topic_model.get_topic(topic_id)}")


for topic_id in topic_info_all['Topic']:
    keywords = [word[0] for word in topic_model.get_topic(topic_id)]
    label = topic_labels.get(topic_id, f"Topic {topic_id}")
    print(f"Topic {topic_id}: {label}")
    print(f"  Top Keywords: {', '.join(keywords)}\n")

topic_model.set_topic_labels(topic_labels)

# %%
import pandas as pd


df_chunks = pd.DataFrame({
    "Interview": chunk_metadata,
    "Chunk": chunks_only,
    "Topic": topics
})


df_chunks = df_chunks[df_chunks["Topic"] != -1]

topic_counts = df_chunks.groupby(["Interview", "Topic"]).size().reset_index(name="Count")


total_counts = df_chunks.groupby("Interview").size().reset_index(name="Total")
topic_counts = topic_counts.merge(total_counts, on="Interview")
topic_counts["Proportion"] = topic_counts["Count"] / topic_counts["Total"]


topic_counts["Label"] = topic_counts["Topic"].map(topic_labels)


print(topic_counts.sort_values(["Interview", "Proportion"], ascending=[True, False]))
print(set(chunk_metadata))


# %%
import matplotlib.pyplot as plt
import seaborn as sns


topic_counts["Interview_ID"] = topic_counts["Interview"].str.extract(r'(\d+)').astype(int) - 1  


topic_counts["Label_With_ID"] = topic_counts["Topic"].astype(str) + ": " + topic_counts["Label"]


plot_df = topic_counts.pivot_table(
    index="Interview_ID",
    columns="Label_With_ID",
    values="Proportion",
    fill_value=0
)


sorted_topic_cols = sorted(plot_df.columns, key=lambda x: int(x.split(":")[0]))
plot_df = plot_df[sorted_topic_cols]
fig, ax = plt.subplots(figsize=(16, 8))
bars = plot_df.plot(
    kind='bar',
    stacked=True,
    colormap='tab20',
    width=0.85,
    ax=ax
)

ax.set_title("Topic Prevalence Across Interviews", fontsize=16)
ax.set_ylabel("Proportion of Chunks", fontsize=12)
ax.set_xlabel("Interview ID (0–12)", fontsize=12)
ax.tick_params(axis='x', labelrotation=0)
ax.tick_params(axis='both', labelsize=10)

ax.legend(
    bbox_to_anchor=(0.5, -0.25),
    loc='upper center',
    ncol=3,
    fontsize=9,
    title='Topic (ID: Label)',
    title_fontsize=10
)

ax.grid(axis='y', linestyle='--', alpha=0.5)
fig.tight_layout()

plt.show()

# %%
for i in range(len(chunks_only)):
    topic_id = topics[i]
    topic_label = topic_labels.get(topic_id, f"Topic {topic_id}")
    topic_keywords = [word[0] for word in topic_model.get_topic(topic_id)]

    print(f"--- Chunk {i + 1} ---")
    print(f"Interview: {chunk_metadata[i]}")
    print(f"Topic ID: {topic_id}")
    print(f"Label: {topic_label}")
    print(f"Keywords: {', '.join(topic_keywords)}")
    print(f"Text: {chunks_only[i][:300]}...")
    print()

# %%

topic_distr, _ = topic_model.approximate_distribution(chunks_only, use_embedding_model=False)

chunk_to_interview = []
for i, text in enumerate(all_cleaned_texts):
    num_chunks = len(list(chunk_sentences(split_text_into_sentences(text), chunk_size)))
    chunk_to_interview.extend([i] * num_chunks) 

soft_df = pd.DataFrame(topic_distr)
soft_df['Interview_ID'] = chunk_to_interview

if -1 in soft_df.columns:
    soft_df = soft_df.drop(columns=[-1])

interview_topic_distr = soft_df.groupby('Interview_ID').mean()

interview_topic_distr = interview_topic_distr[sorted(interview_topic_distr.columns)]
interview_topic_distr = interview_topic_distr.sort_index()


topic_labels_with_id = {}
for topic_id in interview_topic_distr.columns:
    label = topic_labels.get(topic_id, f"Topic {topic_id}")
    topic_labels_with_id[topic_id] = f"{topic_id}: {label}"

interview_topic_distr.rename(columns=topic_labels_with_id, inplace=True)

fig, ax = plt.subplots(figsize=(16, 8))  

interview_topic_distr.plot(
    kind='bar',
    stacked=True,
    colormap='tab20',
    width=0.85,
    ax=ax
)

ax.set_title("Approximate Topic Distribution Per Interview", fontsize=16)
ax.set_ylabel("Average Topic Probability", fontsize=12)
ax.set_xlabel("Interview ID (0–12)", fontsize=12)
ax.tick_params(axis='x', labelrotation=0)
ax.tick_params(axis='both', labelsize=10)

ax.legend(
    bbox_to_anchor=(0.5, -0.25),
    loc='upper center',
    ncol=3,
    fontsize=9,
    title='Topic (ID: Label)',
    title_fontsize=10
)

ax.grid(axis='y', linestyle='--', alpha=0.5)
fig.tight_layout()

plt.show()



# %%
hard_topic_counts = df_chunks.groupby(['Interview', 'Topic']).size().reset_index(name='Count')

top3_hard = hard_topic_counts.sort_values(['Interview', 'Count'], ascending=[True, False]) \
                             .groupby('Interview').head(3)

top3_hard['Label'] = top3_hard['Topic'].map(topic_labels)

print("Top 3 Hard Topics per Interview:")
print(top3_hard)

interview_id_to_name = {i: f"Interview {i+1}" for i in soft_df['Interview_ID'].unique()}
soft_df['Interview'] = soft_df['Interview_ID'].map(interview_id_to_name)

interview_soft_avg = soft_df.drop(columns=['Interview_ID']).groupby('Interview').mean()

top3_soft_list = []
for interview, row in interview_soft_avg.iterrows():
    top3_topics = row.sort_values(ascending=False).head(3)
    for topic_id, prob in top3_topics.items():
        top3_soft_list.append({'Interview': interview, 'Topic': topic_id, 'Avg_Probability': prob})

top3_soft = pd.DataFrame(top3_soft_list)

top3_soft['Label'] = top3_soft['Topic'].map(topic_labels)

print("\nTop 3 Approx. Distrib. Topics per Interview:")
print(top3_soft)

overall_hard = df_chunks['Topic'].value_counts().reset_index()
overall_hard.columns = ['Topic', 'Count']
overall_hard['Label'] = overall_hard['Topic'].map(topic_labels)

print("\nMost Occurring Topics Overall (Hard):")
print(overall_hard.head(5))

overall_soft = interview_soft_avg.mean().sort_values(ascending=False).reset_index()
overall_soft.columns = ['Topic', 'Mean_Avg_Probability']
overall_soft['Label'] = overall_soft['Topic'].map(topic_labels)

print("\nMost Occurring Topics Overall (Approx. Dist.) - Normalized:")
print(overall_soft.head(5))



# %%
